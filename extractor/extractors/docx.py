"""
extractors/docx.py
Extracts text and tables from Microsoft Word (.docx) files.

Install:
    pip install python-docx
"""


def extract_docx(file_path: str) -> list:
    """
    Extracts all paragraphs and tables from a .docx file.
    Returns the full document as a single page dict (Word has no page breaks
    accessible via python-docx).

    Returns:
        List with one dict containing keys:
        page, text, source, type, method, has_tables
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx not installed.\n"
            "Run: pip install python-docx"
        )

    doc    = Document(file_path)
    source = file_path.split("/")[-1]

    # ── Extract all paragraphs ────────────────────────────────
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # ── Extract all tables → pipe-delimited ──────────────────
    table_text = ""
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_text += " | ".join(cells) + "\n"

    # ── Combine ───────────────────────────────────────────────
    full_text = "\n".join(paragraphs)
    if table_text.strip():
        full_text += "\n\n[TABLE DATA]\n" + table_text

    return [{
        "page":       1,
        "text":       full_text.strip(),
        "source":     source,
        "type":       "docx",
        "method":     "python-docx",
        "has_tables": bool(doc.tables),
    }]